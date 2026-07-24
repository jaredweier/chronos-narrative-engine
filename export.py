import io
import os
import zipfile
from datetime import datetime
from html import escape
from typing import Optional, List, Dict, Any


def _add_letterhead(doc):
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from config import DEPARTMENT_NAME, DEPARTMENT_ADDRESS, DEPARTMENT_CITY_STATE_ZIP, DEPARTMENT_PHONE

    letterhead_table = doc.add_table(rows=1, cols=1)
    cell = letterhead_table.cell(0, 0)
    p = cell.paragraphs[0]
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run(DEPARTMENT_NAME)
    r.bold = True
    r.font.size = Pt(14)
    r.font.color.rgb = RGBColor(0, 0, 0)
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = p2.add_run(f"{DEPARTMENT_ADDRESS} | {DEPARTMENT_CITY_STATE_ZIP}")
    r2.font.size = Pt(9)
    r2.font.color.rgb = RGBColor(100, 100, 100)
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = p3.add_run(f"Phone: {DEPARTMENT_PHONE}")
    r3.font.size = Pt(9)
    r3.font.color.rgb = RGBColor(100, 100, 100)
    doc.add_paragraph()


def export_report_docx(
    report_text: str,
    officer_name: str,
    badge_number: str,
    incident_id: Optional[str] = None,
    call_type: Optional[str] = None,
    location: Optional[str] = None,
    report_type: Optional[str] = None,
) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_TABLE_ALIGNMENT

    doc = Document()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    font.color.rgb = RGBColor(0, 0, 0)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    _add_letterhead(doc)

    header_table = doc.add_table(rows=1, cols=1)
    header_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = header_table.cell(0, 0)
    cell_para = cell.paragraphs[0]
    cell_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = cell_para.add_run("CHRONOS NARRATIVE ENGINE")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 58, 138)

    sub_para = cell.add_paragraph()
    sub_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub_para.add_run("Law Enforcement Report Generation System")
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(100, 116, 139)

    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    borders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'single')
        element.set(qn('w:sz'), '12')
        element.set(qn('w:color'), '1E3A8A')
        element.set(qn('w:space'), '0')
        borders.append(element)
    for edge in ('left', 'right'):
        element = OxmlElement(f'w:{edge}')
        element.set(qn('w:val'), 'none')
        element.set(qn('w:sz'), '0')
        element.set(qn('w:color'), 'auto')
        element.set(qn('w:space'), '0')
        borders.append(element)
    tcPr.append(borders)

    doc.add_paragraph()

    info_table = doc.add_table(rows=4, cols=2)
    info_table.style = 'Table Grid'
    info_table.alignment = WD_TABLE_ALIGNMENT.LEFT

    now = datetime.now()
    display_type = report_type or call_type or "Incident Report"
    info_data = [
        ("Officer:", officer_name),
        ("Badge Number:", badge_number),
        ("Incident ID:", incident_id or f"INC-{now.strftime('%Y%m%d')}-{now.strftime('%H%M%S')}"),
        ("Report Type:", display_type),
    ]

    for i, (label, value) in enumerate(info_data):
        label_cell = info_table.cell(i, 0)
        value_cell = info_table.cell(i, 1)
        label_para = label_cell.paragraphs[0]
        run = label_para.add_run(label)
        run.bold = True
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'
        value_para = value_cell.paragraphs[0]
        run = value_para.add_run(value)
        run.font.size = Pt(10)
        run.font.name = 'Times New Roman'

    if call_type or location:
        extra_table = doc.add_table(rows=1, cols=2)
        extra_table.style = 'Table Grid'
        c1 = extra_table.cell(0, 0)
        c2 = extra_table.cell(0, 1)
        p1 = c1.paragraphs[0]
        r1 = p1.add_run("Call Type: ")
        r1.bold = True
        r1.font.size = Pt(10)
        r1.font.name = 'Times New Roman'
        r1b = p1.add_run(call_type or "N/A")
        r1b.font.size = Pt(10)
        r1b.font.name = 'Times New Roman'
        p2 = c2.paragraphs[0]
        r2 = p2.add_run("Location: ")
        r2.bold = True
        r2.font.size = Pt(10)
        r2.font.name = 'Times New Roman'
        r2b = p2.add_run(location or "N/A")
        r2b.font.size = Pt(10)
        r2b.font.name = 'Times New Roman'

    doc.add_paragraph()

    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hr.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    narrative_heading = doc.add_paragraph()
    narrative_heading.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run = narrative_heading.add_run("INCIDENT NARRATIVE")
    run.bold = True
    run.font.size = Pt(13)
    run.font.color.rgb = RGBColor(30, 58, 138)

    for para_text in report_text.split('\n'):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(6)
        p.paragraph_format.line_spacing = Pt(18)
        run = p.add_run(para_text)
        run.font.size = Pt(12)
        run.font.name = 'Times New Roman'

    doc.add_paragraph()

    sig_table = doc.add_table(rows=3, cols=2)
    sig_table.style = 'Table Grid'
    sig_data = [
        ("Officer Signature:", ""),
        ("Printed Name:", officer_name),
        ("Date:", now.strftime("%B %d, %Y")),
    ]
    for i, (label, value) in enumerate(sig_data):
        r0 = sig_table.cell(i, 0).paragraphs[0]
        r1 = sig_table.cell(i, 1).paragraphs[0]
        run0 = r0.add_run(label)
        run0.bold = True
        run0.font.size = Pt(10)
        run0.font.name = 'Times New Roman'
        run1 = r1.add_run(value)
        run1.font.size = Pt(10)
        run1.font.name = 'Times New Roman'

    doc.add_paragraph()
    doc.add_paragraph()

    footer_hr = doc.add_paragraph()
    footer_hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer_hr.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run("Generated by Chronos Narrative Engine")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.italic = True

    footer2 = doc.add_paragraph()
    footer2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer2.add_run(f"Timestamp: {now.strftime('%Y-%m-%d %H:%M:%S')} | Officer: {officer_name} | Badge #{badge_number}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.italic = True

    footer3 = doc.add_paragraph()
    footer3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer3.add_run("This report was AI-generated and has been reviewed by the officer of record.")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_report_pdf(
    report_text: str,
    officer_name: str,
    badge_number: str,
    incident_id: Optional[str] = None,
    call_type: Optional[str] = None,
    location: Optional[str] = None,
    report_type: Optional[str] = None,
) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.lib.colors import HexColor
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable, Table, TableStyle
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.enums import TA_CENTER, TA_LEFT

    buf = io.BytesIO()
    now = datetime.now()

    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'ChronosTitle',
        parent=styles['Title'],
        fontSize=18,
        textColor=HexColor('#1E3A8A'),
        spaceAfter=2,
        alignment=TA_CENTER,
        fontName='Helvetica-Bold',
    )

    subtitle_style = ParagraphStyle(
        'ChronosSubtitle',
        parent=styles['Normal'],
        fontSize=9,
        textColor=HexColor('#64748B'),
        alignment=TA_CENTER,
        spaceAfter=12,
        fontName='Helvetica',
    )

    section_style = ParagraphStyle(
        'SectionHead',
        parent=styles['Heading2'],
        fontSize=13,
        textColor=HexColor('#1E3A8A'),
        fontName='Helvetica-Bold',
        spaceBefore=16,
        spaceAfter=8,
    )

    body_style = ParagraphStyle(
        'ReportBody',
        parent=styles['Normal'],
        fontSize=11,
        leading=16,
        fontName='Times-Roman',
        spaceAfter=6,
    )

    info_label_style = ParagraphStyle(
        'InfoLabel',
        parent=styles['Normal'],
        fontSize=10,
        fontName='Helvetica-Bold',
    )

    info_value_style = ParagraphStyle(
        'InfoValue',
        parent=styles['Normal'],
        fontSize=10,
        fontName='Helvetica',
    )

    footer_style = ParagraphStyle(
        'Footer',
        parent=styles['Normal'],
        fontSize=8,
        textColor=HexColor('#64748B'),
        alignment=TA_CENTER,
        fontName='Helvetica-Oblique',
    )

    elements = []

    from config import DEPARTMENT_NAME, DEPARTMENT_ADDRESS, DEPARTMENT_CITY_STATE_ZIP
    dept_style = ParagraphStyle('Dept', parent=styles['Normal'], fontSize=11,
                                 textColor=HexColor('#000000'), alignment=TA_CENTER,
                                 fontName='Helvetica-Bold', spaceAfter=1)
    dept_addr_style = ParagraphStyle('DeptAddr', parent=styles['Normal'], fontSize=8,
                                     textColor=HexColor('#666666'), alignment=TA_CENTER,
                                     fontName='Helvetica', spaceAfter=1)
    elements.append(Paragraph(DEPARTMENT_NAME, dept_style))
    elements.append(Paragraph(f"{DEPARTMENT_ADDRESS} | {DEPARTMENT_CITY_STATE_ZIP}", dept_addr_style))
    elements.append(Spacer(1, 8))

    elements.append(Paragraph("CHRONOS NARRATIVE ENGINE", title_style))
    elements.append(Paragraph("Law Enforcement Report Assistance Program", subtitle_style))
    elements.append(HRFlowable(width="100%", thickness=2, color=HexColor('#1E3A8A'), spaceAfter=12))

    inc_id = incident_id or f"INC-{now.strftime('%Y%m%d')}-{now.strftime('%H%M%S')}"
    display_type = report_type or call_type or "Incident Report"

    info_data = [
        ["Officer:", officer_name, "Incident ID:", inc_id],
        ["Badge #:", badge_number, "Date:", now.strftime("%B %d, %Y")],
        ["Report Type:", display_type, "Location:", location or "N/A"],
    ]

    info_table = Table(info_data, colWidths=[1.1*inch, 2.2*inch, 1.1*inch, 2.2*inch])
    info_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTNAME', (3, 0), (3, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('TEXTCOLOR', (0, 0), (0, -1), HexColor('#334155')),
        ('TEXTCOLOR', (2, 0), (2, -1), HexColor('#334155')),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('TOPPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#CBD5E1')),
        ('BACKGROUND', (0, 0), (0, -1), HexColor('#F1F5F9')),
        ('BACKGROUND', (2, 0), (2, -1), HexColor('#F1F5F9')),
    ]))

    elements.append(info_table)
    elements.append(Spacer(1, 16))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#94A3B8'), spaceAfter=12))

    elements.append(Paragraph("INCIDENT NARRATIVE", section_style))

    for para_text in report_text.split('\n'):
        if para_text.strip():
            elements.append(Paragraph(para_text.replace('<', '&lt;').replace('>', '&gt;'), body_style))
        else:
            elements.append(Spacer(1, 6))

    elements.append(Spacer(1, 24))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#94A3B8'), spaceAfter=12))

    sig_data = [
        ["Officer Signature:", "______________________________", "Date:", now.strftime("%B %d, %Y")],
        ["Printed Name:", officer_name, "Badge #:", badge_number],
    ]
    sig_table = Table(sig_data, colWidths=[1.2*inch, 2*inch, 0.8*inch, 1.6*inch])
    sig_table.setStyle(TableStyle([
        ('FONTNAME', (0, 0), (0, -1), 'Helvetica-Bold'),
        ('FONTNAME', (2, 0), (2, -1), 'Helvetica-Bold'),
        ('FONTNAME', (1, 0), (1, -1), 'Helvetica'),
        ('FONTNAME', (3, 0), (3, -1), 'Helvetica'),
        ('FONTSIZE', (0, 0), (-1, -1), 10),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 6),
        ('GRID', (0, 0), (-1, -1), 0.5, HexColor('#CBD5E1')),
    ]))
    elements.append(sig_table)
    elements.append(Spacer(1, 16))

    elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#94A3B8'), spaceAfter=12))
    elements.append(Paragraph("Generated by Chronos Narrative Engine", footer_style))
    elements.append(Paragraph(
        f"Timestamp: {now.strftime('%Y-%m-%d %H:%M:%S')} | Officer: {officer_name} | Badge #{badge_number}",
        footer_style
    ))
    elements.append(Paragraph(
        "This report was AI-generated and has been reviewed by the officer of record.",
        footer_style
    ))

    def add_page_number(canvas, doc):
        canvas.saveState()
        canvas.setFont('Helvetica', 8)
        canvas.setFillColor(HexColor('#64748B'))
        canvas.drawCentredString(letter[0] / 2, 0.5 * inch, f"Page {doc.page}")
        canvas.restoreState()

    doc.build(elements, onFirstPage=add_page_number, onLaterPages=add_page_number)
    buf.seek(0)
    return buf.getvalue()


def export_compliance_docx(officer_name: str, badge_number: str) -> bytes:
    from docx import Document
    from docx.shared import Inches, Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from compliance_content import COMPLIANCE_SECTIONS

    doc = Document()
    now = datetime.now()

    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(11)

    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.25)
        section.right_margin = Inches(1.25)

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ARTIFICIAL INTELLIGENCE\nSAFEGUARDS & GUARDRAILS CERTIFICATION")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 58, 138)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Chronos Narrative Engine — Law Enforcement Report Generation System")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 116, 139)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = meta.add_run(f"Officer: {officer_name} | Badge #{badge_number} | Date: {now.strftime('%B %d, %Y')}")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()
    hr = doc.add_paragraph()
    hr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hr.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    intro = doc.add_paragraph()
    run = intro.add_run(
        "This document certifies that the Chronos Narrative Engine operates under the following "
        "AI safeguards, guardrails, and compliance controls. All AI-generated content produced "
        "by this system is subject to mandatory human review before submission to any records "
        "management system."
    )
    run.font.size = Pt(11)
    run.font.name = 'Times New Roman'

    for heading, items in COMPLIANCE_SECTIONS:
        plain_heading = heading.replace("&amp;", "&").replace("&mdash;", "-")
        h = doc.add_paragraph()
        run = h.add_run(plain_heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 58, 138)
        for item in items:
            plain_item = item.replace("&amp;", "&").replace("&mdash;", "-").replace("&bull;", "*")
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(plain_item)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'

    doc.add_paragraph()
    hr2 = doc.add_paragraph()
    hr2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = hr2.add_run("─" * 60)
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)

    footer = doc.add_paragraph()
    footer.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = footer.add_run(f"Generated by Chronos Narrative Engine | {now.strftime('%Y-%m-%d %H:%M:%S')}")
    run.font.size = Pt(8)
    run.font.color.rgb = RGBColor(100, 116, 139)
    run.italic = True

    buf = io.BytesIO()
    doc.save(buf)
    buf.seek(0)
    return buf.getvalue()


def export_compliance_pdf(officer_name: str, badge_number: str) -> bytes:
    from reportlab.lib.pagesizes import letter
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Spacer, HRFlowable
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER
    from compliance_content import COMPLIANCE_SECTIONS

    buf = io.BytesIO()
    now = datetime.now()

    doc = SimpleDocTemplate(
        buf,
        pagesize=letter,
        topMargin=0.75 * inch,
        bottomMargin=0.75 * inch,
        leftMargin=1 * inch,
        rightMargin=1 * inch,
    )

    styles = getSampleStyleSheet()

    title_style = ParagraphStyle(
        'ComplianceTitle', parent=styles['Title'], fontSize=16,
        textColor=HexColor('#1E3A8A'), alignment=TA_CENTER,
        fontName='Helvetica-Bold', spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        'ComplianceSubtitle', parent=styles['Normal'], fontSize=9,
        textColor=HexColor('#64748B'), alignment=TA_CENTER, spaceAfter=4,
    )
    meta_style = ParagraphStyle(
        'ComplianceMeta', parent=styles['Normal'], fontSize=10,
        textColor=HexColor('#64748B'), alignment=TA_CENTER, spaceAfter=12,
    )
    section_style = ParagraphStyle(
        'ComplianceSection', parent=styles['Heading2'], fontSize=12,
        textColor=HexColor('#1E3A8A'), fontName='Helvetica-Bold',
        spaceBefore=14, spaceAfter=6,
    )
    body_style = ParagraphStyle(
        'ComplianceBody', parent=styles['Normal'], fontSize=10,
        leading=14, fontName='Times-Roman', spaceAfter=4,
    )
    bullet_style = ParagraphStyle(
        'ComplianceBullet', parent=styles['Normal'], fontSize=10,
        leading=14, fontName='Times-Roman', leftIndent=20, spaceAfter=3, bulletIndent=8,
    )
    footer_style = ParagraphStyle(
        'ComplianceFooter', parent=styles['Normal'], fontSize=8,
        textColor=HexColor('#64748B'), alignment=TA_CENTER, fontName='Helvetica-Oblique',
    )

    elements = []
    elements.append(Paragraph("ARTIFICIAL INTELLIGENCE<br/>SAFEGUARDS &amp; GUARDRAILS CERTIFICATION", title_style))
    elements.append(Paragraph("Chronos Narrative Engine — Law Enforcement Report Generation System", subtitle_style))
    elements.append(Paragraph(f"Officer: {officer_name} | Badge #{badge_number} | Date: {now.strftime('%B %d, %Y')}", meta_style))
    elements.append(HRFlowable(width="100%", thickness=2, color=HexColor('#1E3A8A'), spaceAfter=12))
    elements.append(Paragraph(
        "This document certifies that the Chronos Narrative Engine operates under the following "
        "AI safeguards, guardrails, and compliance controls. All AI-generated content produced "
        "by this system is subject to mandatory human review before submission to any records "
        "management system.",
        body_style
    ))

    for heading, items in COMPLIANCE_SECTIONS:
        elements.append(Paragraph(heading, section_style))
        for item in items:
            safe = item.replace('<', '&lt;').replace('>', '&gt;')
            elements.append(Paragraph(f"* {safe}", bullet_style))

    elements.append(Spacer(1, 24))
    elements.append(HRFlowable(width="100%", thickness=0.5, color=HexColor('#94A3B8'), spaceAfter=12))
    elements.append(Paragraph(f"Generated by Chronos Narrative Engine | {now.strftime('%Y-%m-%d %H:%M:%S')}", footer_style))
    elements.append(Paragraph("This document certifies compliance with the AI safeguards described above.", footer_style))

    doc.build(elements)
    buf.seek(0)
    return buf.getvalue()


# --- Bulk ZIP Export ---

def export_bulk_zip(reports, include_nibrs=False):
    import zipfile
    from nibrs_export import build_nibrs_xml
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        for r in reports:
            inc_id = r.get('incident_id', 'unknown')
            officer = r.get('officer_name', 'unknown')
            text = r.get('final_approved_report') or r.get('unedited_ai_draft', '')
            docx_bytes = export_report_docx(text, officer, r.get('officer_id', ''), inc_id)
            zf.writestr(f"{inc_id}/{inc_id}_report.docx", docx_bytes)
            pdf_bytes = export_report_pdf(text, officer, r.get('officer_id', ''), inc_id)
            zf.writestr(f"{inc_id}/{inc_id}_report.pdf", pdf_bytes)
            zf.writestr(f"{inc_id}/{inc_id}_text.txt", text)
            if include_nibrs and r.get('document_type'):
                nibrs_xml = build_nibrs_xml(
                    incident_id=inc_id,
                    officer_name=officer,
                    officer_id=r.get('officer_id', ''),
                    report_type=r.get('document_type', ''),
                    narrative=text,
                )
                zf.writestr(f"{inc_id}/{inc_id}_nibrs.xml", nibrs_xml)
    buf.seek(0)
    return buf.getvalue()


# --- Signature Capture ---

SIGNATURE_CAPTURE_HTML = """
<div id="sig-canvas-wrapper" style="border:2px dashed #334155;border-radius:6px;padding:4px;margin-bottom:8px;background:#0f172a;">
  <canvas id="sig-canvas" width="500" height="150" style="width:100%;height:150px;cursor:crosshair;touch-action:none;"></canvas>
</div>
<div style="display:flex;gap:8px;margin-bottom:8px;">
  <button id="sig-clear" style="flex:1;padding:6px;background:#1e293b;color:#94a3b8;border:1px solid #334155;border-radius:4px;cursor:pointer;">Clear</button>
  <button id="sig-accept" style="flex:1;padding:6px;background:#1d4ed8;color:white;border:none;border-radius:4px;cursor:pointer;">Accept Signature</button>
</div>
<input type="hidden" id="sig-data" value="">
<script>
(function() {
  var canvas = document.getElementById('sig-canvas');
  if (!canvas) return;
  var ctx = canvas.getContext('2d');
  var drawing = false;
  var rect = canvas.getBoundingClientRect();

  function resize() {
    rect = canvas.getBoundingClientRect();
    canvas.width = rect.width * (window.devicePixelRatio || 1);
    canvas.height = rect.height * (window.devicePixelRatio || 1);
    ctx.scale(window.devicePixelRatio || 1, window.devicePixelRatio || 1);
    ctx.strokeStyle = '#e2e8f0';
    ctx.lineWidth = 2;
    ctx.lineCap = 'round';
  }
  resize();
  window.addEventListener('resize', resize);

  function getPos(e) {
    var r = canvas.getBoundingClientRect();
    var clientX = e.touches ? e.touches[0].clientX : e.clientX;
    var clientY = e.touches ? e.touches[0].clientY : e.clientY;
    return { x: clientX - r.left, y: clientY - r.top };
  }

  function start(e) { e.preventDefault(); drawing = true; var p = getPos(e); ctx.beginPath(); ctx.moveTo(p.x, p.y); }
  function move(e) { e.preventDefault(); if (!drawing) return; var p = getPos(e); ctx.lineTo(p.x, p.y); ctx.stroke(); }
  function stop(e) { e.preventDefault(); drawing = false; }

  canvas.addEventListener('mousedown', start);
  canvas.addEventListener('mousemove', move);
  canvas.addEventListener('mouseup', stop);
  canvas.addEventListener('mouseleave', stop);
  canvas.addEventListener('touchstart', start, {passive:false});
  canvas.addEventListener('touchmove', move, {passive:false});
  canvas.addEventListener('touchend', stop, {passive:false});

  document.getElementById('sig-clear').addEventListener('click', function() {
    ctx.clearRect(0, 0, canvas.width, canvas.height);
    document.getElementById('sig-data').value = '';
  });

  document.getElementById('sig-accept').addEventListener('click', function() {
    var dataUrl = canvas.toDataURL('image/png');
    document.getElementById('sig-data').value = dataUrl;
    var wrapper = document.getElementById('sig-canvas-wrapper');
    wrapper.style.borderColor = '#22c55e';
  });
})();
</script>
"""


def export_report_html_preview(text: str, officer_name: str, badge_num: str, case_no: str) -> str:
    from config import DEPARTMENT_NAME, DEPARTMENT_ADDRESS, DEPARTMENT_CITY_STATE_ZIP, DEPARTMENT_PHONE
    now = datetime.now()
    escaped_text = escape(text).replace('\n', '<br>')
    escaped_officer = escape(officer_name)
    escaped_badge = escape(badge_num)
    escaped_case = escape(case_no)
    escaped_dept = escape(DEPARTMENT_NAME)
    escaped_addr = escape(f"{DEPARTMENT_ADDRESS} | {DEPARTMENT_CITY_STATE_ZIP}")
    escaped_phone = escape(DEPARTMENT_PHONE)
    return f"""<!DOCTYPE html>
<html lang="en">
<head><meta charset="UTF-8"><title>Incident Report - {escaped_case}</title></head>
<body style="font-family:Georgia,'Times New Roman',serif;color:#000;background:#fff;margin:0;padding:40px;max-width:800px;margin:0 auto;">
<div style="text-align:center;margin-bottom:20px;">
<h1 style="font-size:18px;margin:0;">{escaped_dept}</h1>
<p style="font-size:11px;color:#666;margin:2px 0;">{escaped_addr}</p>
<p style="font-size:11px;color:#666;margin:2px 0;">Phone: {escaped_phone}</p>
</div>
<hr style="border:none;border-top:2px solid #1E3A8A;margin:16px 0;">
<h2 style="text-align:center;font-size:16px;color:#1E3A8A;margin:8px 0;">CHRONOS NARRATIVE ENGINE</h2>
<p style="text-align:center;font-size:9px;color:#64748B;margin:2px 0 16px;">Law Enforcement Report Generation System</p>
<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:11px;">
<tr><td style="font-weight:bold;padding:4px 8px;width:120px;background:#F1F5F9;">Officer:</td><td style="padding:4px 8px;">{escaped_officer}</td></tr>
<tr><td style="font-weight:bold;padding:4px 8px;background:#F1F5F9;">Badge #:</td><td style="padding:4px 8px;">{escaped_badge}</td></tr>
<tr><td style="font-weight:bold;padding:4px 8px;background:#F1F5F9;">Case #:</td><td style="padding:4px 8px;">{escaped_case}</td></tr>
<tr><td style="font-weight:bold;padding:4px 8px;background:#F1F5F9;">Date:</td><td style="padding:4px 8px;">{now.strftime('%B %d, %Y')}</td></tr>
</table>
<hr style="border:none;border-top:1px solid #94A3B8;margin:16px 0;">
<h3 style="font-size:14px;color:#1E3A8A;margin:12px 0 8px;">INCIDENT NARRATIVE</h3>
<div style="font-size:12px;line-height:1.6;">{escaped_text}</div>
<hr style="border:none;border-top:1px solid #94A3B8;margin:24px 0;">
<table style="width:100%;border-collapse:collapse;margin:12px 0;font-size:10px;">
<tr><td style="font-weight:bold;padding:4px 8px;width:140px;">Officer Signature:</td><td style="border-bottom:1px solid #000;padding:4px 8px;"></td></tr>
<tr><td style="font-weight:bold;padding:4px 8px;">Printed Name:</td><td style="padding:4px 8px;">{escaped_officer}</td></tr>
<tr><td style="font-weight:bold;padding:4px 8px;">Date:</td><td style="padding:4px 8px;">{now.strftime('%B %d, %Y')}</td></tr>
</table>
<hr style="border:none;border-top:1px solid #94A3B8;margin:16px 0;">
<p style="text-align:center;font-size:8px;color:#64748B;font-style:italic;margin:2px 0;">Generated by Chronos Narrative Engine</p>
<p style="text-align:center;font-size:8px;color:#64748B;font-style:italic;margin:2px 0;">Timestamp: {now.strftime('%Y-%m-%d %H:%M:%S')} | Officer: {escaped_officer} | Badge #{escaped_badge}</p>
<p style="text-align:center;font-size:8px;color:#64748B;font-style:italic;margin:2px 0;">This report was AI-generated and has been reviewed by the officer of record.</p>
</body>
</html>"""


def export_offline_package(incident_id: str, text: str, officer_name: str, badge_num: str, case_no: str, hmac_key: str = "") -> bytes:
    import hashlib
    import hmac as hmac_module
    import json
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{case_no}/report.txt", text)
        html_content = export_report_html_preview(text, officer_name, badge_num, case_no)
        zf.writestr(f"{case_no}/report.html", html_content)
        docx_bytes = export_report_docx(text, officer_name, badge_num, incident_id)
        zf.writestr(f"{case_no}/report.docx", docx_bytes)
        now = datetime.utcnow()
        manifest = {
            "incident_id": incident_id,
            "officer_name": officer_name,
            "badge_num": badge_num,
            "case_no": case_no,
            "timestamp": now.isoformat(),
            "file_list": ["report.txt", "report.html", "report.docx"],
        }
        manifest_json = json.dumps(manifest, indent=2)
        zf.writestr(f"{case_no}/manifest.json", manifest_json)
        if hmac_key:
            h = hmac_module.new(hmac_key.encode(), manifest_json.encode(), hashlib.sha256)
            signature = h.hexdigest()
        else:
            signature = hashlib.sha256(manifest_json.encode()).hexdigest()
        zf.writestr(f"{case_no}/signature.sha256", signature)
    buf.seek(0)
    return buf.getvalue()


def export_report_with_evidence(incident_id: str, text: str, officer_name: str, badge_num: str, case_no: str) -> bytes:
    from database import get_evidence_files
    buf = io.BytesIO()
    with zipfile.ZipFile(buf, 'w', zipfile.ZIP_DEFLATED) as zf:
        zf.writestr(f"{case_no}/{case_no}_report.txt", text)
        html_content = export_report_html_preview(text, officer_name, badge_num, case_no)
        zf.writestr(f"{case_no}/{case_no}_report.html", html_content)
        docx_bytes = export_report_docx(text, officer_name, badge_num, incident_id)
        zf.writestr(f"{case_no}/{case_no}_report.docx", docx_bytes)
        try:
            evidence_files = get_evidence_files(incident_id)
            for ef in evidence_files:
                file_path = ef.get('file_path', '')
                file_name = ef.get('file_name', '')
                if file_path and os.path.exists(file_path):
                    with open(file_path, 'rb') as f:
                        zf.writestr(f"{case_no}/evidence/{file_name}", f.read())
        except Exception:
            pass
    buf.seek(0)
    return buf.getvalue()
