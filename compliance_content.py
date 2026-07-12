from html import escape as h
from datetime import datetime


COMPLIANCE_SECTIONS = [
    ("1. Purpose &amp; Scope", [
        "The Chronos Narrative Engine is designed to assist law enforcement officers in drafting incident reports and related documentation",
        "This system is an aid tool &mdash; it does not replace officer judgment, observation, or professional responsibility",
        "All generated content constitutes a preliminary draft requiring thorough officer review before any official submission",
        "The system is intended for use by trained law enforcement personnel only",
    ]),
    ("2. Data Sovereignty &amp; Local Operation", [
        "All inference is performed locally via Ollama (localhost:11434) &mdash; no data leaves the authorized workstation",
        "No cloud-based AI services, APIs, or external endpoints are contacted during any processing stage",
        "All transcriptions, extractions, and narrative generation occur on-device using the local LLM (Llama 3.1:8b)",
        "CJIS compliance is maintained through complete data sovereignty on the local machine",
        "No data is transmitted to, stored by, or accessible from any third-party service at any time",
    ]),
    ("3. PII Protection &amp; Redaction", [
        "All personally identifiable information (SSNs, DOBs, phone numbers, emails, juvenile names) is automatically detected and redacted before entering any AI context",
        "Redaction is irreversible within the AI pipeline &mdash; original data is never transmitted to the LLM",
        "A full redaction audit trail is maintained in the SQLite database for every processed report",
        "The redactor supports multiple detection patterns including SSNs, phone numbers, email addresses, and juvenile identifiers",
        "Officers should verify redaction completeness before and after AI processing",
    ]),
    ("4. Human-in-the-Loop Requirements", [
        "All AI-generated narratives are presented in an editable text area before submission",
        "Officers must explicitly review and may modify any AI-generated content before submission",
        "The system logs the original AI draft alongside any human edits for audit purposes",
        "No report is automatically submitted without officer verification and explicit action",
        "The officer of record bears full legal responsibility for the accuracy and completeness of every submitted report",
    ]),
    ("5. NIBRS Compliance Validation", [
        "Every generated report undergoes automated NIBRS (National Incident-Based Reporting System) compliance checking",
        "Critical compliance violations are flagged and must be resolved before submission to records management",
        "Compliance warnings are presented with severity levels: critical, warning, and informational",
        "Officers should address all critical and warning-level items prior to finalizing any report",
    ]),
    ("6. Audit Trail &amp; Accountability", [
        "Every submission is logged with officer name, badge number, timestamp, case number, and report content",
        "The system tracks AI-generated drafts versus human-edited final versions for accountability",
        "A corrections log is maintained for post-submission amendments and amendments history",
        "All audit data is stored in the local SQLite database (department_reports.db) on the authorized workstation",
        "Audit records are immutable once created and cannot be altered through the application interface",
    ]),
    ("7. Model Limitations &amp; Disclosure", [
        "The AI model (Llama 3.1:8b) is a large language model and does not possess legal expertise, law enforcement training, or investigative judgment",
        "Generated narratives are drafting assistance only &mdash; they require officer review, verification of facts, and domain-specific knowledge",
        "The system does not make legal determinations, probable cause assessments, charging decisions, or any authoritative conclusions",
        "Officers are solely responsible for the accuracy, completeness, legal sufficiency, and truthfulness of all submitted reports",
        "The model may occasionally generate plausible-sounding but inaccurate statements &mdash; all content must be independently verified",
    ]),
    ("8. Access Controls &amp; Authentication", [
        "System access requires officer authentication (name, badge number, and password) at each session",
        "Session state is cleared upon sign-out to prevent unauthorized access to in-progress work",
        "All processing occurs on a single authorized workstation in a controlled environment",
        "Officer credentials are stored using salted SHA-256 hashes",
    ]),
    ("9. Officer Style Matching", [
        "Few-shot style matching is used to match individual officer writing patterns based on previously submitted reports",
        "Style samples are officer-provided redacted reports &mdash; no synthetic, fabricated, or third-party data is used",
        "Officers retain full control over which samples are used for style training and may remove samples at any time",
        "Style matching influences tone and structure only &mdash; factual content is derived exclusively from evidence and officer input",
    ]),
    ("10. Evidence Handling", [
        "CAD report data and body camera transcripts are processed locally and never transmitted externally",
        "Uploaded evidence files are temporarily stored in the processing directory and removed after extraction",
        "Evidence data is redacted before being used in any AI context to protect sensitive information",
        "Officers should verify that all evidence has been properly accounted for after processing",
    ]),
    ("11. Limitations of Liability", [
        "This system is provided as a drafting assistance tool &mdash; it does not guarantee accuracy, completeness, or legal sufficiency of any generated content",
        "The department and system administrators are not liable for errors in AI-generated draft content",
        "Officers assume full professional and legal responsibility for all reports submitted under their name and badge number",
        "This system has not been certified by any external authority and should be used in conjunction with department policies and procedures",
    ]),
]


def get_compliance_html(officer: str = "", badge_num: str = "") -> str:
    today = datetime.now().strftime("%B %d, %Y")
    sections_html = ""
    for heading, items in COMPLIANCE_SECTIONS:
        lis = "".join(f"<li>{item}</li>" for item in items)
        sections_html += f"<h4>{heading}</h4>\n<ul>{lis}</ul>\n"
    officer_h = h(officer) if officer else ""
    badge_h = h(badge_num) if badge_num else ""
    return f"""<div class="compliance-wrap"><div class="compliance-paper">
    <div class="cp-header">
    <h2>Artificial Intelligence<br>Safeguards &amp; Guardrails Certification</h2>
    <div class="cp-sub">Chronos Narrative Engine &mdash; Law Enforcement Report Assistance Program</div>
    <div class="cp-meta">Officer: {officer_h} &nbsp;|&nbsp; Badge #{badge_h} &nbsp;|&nbsp; Date: {today}</div></div>
    <p>This document certifies that the Chronos Narrative Engine operates under the following AI safeguards, guardrails, and compliance controls. All AI-generated content produced by this system is subject to mandatory human review before submission to any records management system.</p>
    {sections_html}
    <div class="cp-footer">Generated by Chronos Narrative Engine &bull; {today}<br>This document certifies compliance with the AI safeguards described above.</div>
    </div></div>"""


def get_compliance_text(officer: str = "", badge_num: str = "") -> str:
    today = datetime.now().strftime("%B %d, %Y")
    lines = [
        "ARTIFICIAL INTELLIGENCE SAFEGUARDS & GUARDRAILS CERTIFICATION",
        "Chronos Narrative Engine - Law Enforcement Report Assistance Program",
        f"Generated: {today}",
        f"Officer: {officer} | Badge #{badge_num}",
        "",
    ]
    for heading, items in COMPLIANCE_SECTIONS:
        plain_heading = heading.replace("&amp;", "&").replace("&mdash;", "-")
        lines.append(plain_heading)
        for item in items:
            plain_item = item.replace("&amp;", "&").replace("&mdash;", "-").replace("&bull;", "*")
            lines.append(f"- {plain_item}")
        lines.append("")
    lines.append("This document certifies compliance with the AI safeguards described above.")
    return "\n".join(lines)


def get_compliance_docx_sections(doc):
    from docx.shared import Pt, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("ARTIFICIAL INTELLIGENCE\nSAFEGUARDS & GUARDRAILS CERTIFICATION")
    run.bold = True
    run.font.size = Pt(16)
    run.font.color.rgb = RGBColor(30, 58, 138)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = sub.add_run("Chronos Narrative Engine - Law Enforcement Report Generation System")
    run.font.size = Pt(10)
    run.font.color.rgb = RGBColor(100, 116, 139)

    doc.add_paragraph()

    for heading, items in COMPLIANCE_SECTIONS:
        plain_heading = heading.replace("&amp;", "&").replace("&mdash;", "-")
        h = doc.add_paragraph()
        run = h.add_run(plain_heading)
        run.bold = True
        run.font.size = Pt(12)
        run.font.color.rgb = RGBColor(30, 58, 138)
        for item in items:
            plain_item = item.replace("&amp;", "&").replace("&mdash;", "-").replace("&bull;", "*")
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(plain_item)
            run.font.size = Pt(10)
            run.font.name = 'Times New Roman'


def get_compliance_pdf_elements():
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.colors import HexColor
    from reportlab.lib.enums import TA_CENTER
    from reportlab.platypus import Paragraph, Spacer, HRFlowable

    styles = getSampleStyleSheet()
    elements = []

    title_style = ParagraphStyle(
        'ComplianceTitle', parent=styles['Title'], fontSize=16,
        textColor=HexColor('#1E3A8A'), alignment=TA_CENTER,
        fontName='Helvetica-Bold', spaceAfter=4,
    )
    subtitle_style = ParagraphStyle(
        'ComplianceSubtitle', parent=styles['Normal'], fontSize=9,
        textColor=HexColor('#64748B'), alignment=TA_CENTER, spaceAfter=12,
    )
    section_style = ParagraphStyle(
        'ComplianceSection', parent=styles['Heading2'], fontSize=12,
        textColor=HexColor('#1E3A8A'), fontName='Helvetica-Bold',
        spaceBefore=14, spaceAfter=6,
    )
    bullet_style = ParagraphStyle(
        'ComplianceBullet', parent=styles['Normal'], fontSize=10,
        leading=14, fontName='Times-Roman', leftIndent=20, spaceAfter=3, bulletIndent=8,
    )

    elements.append(Paragraph("ARTIFICIAL INTELLIGENCE<br/>SAFEGUARDS &amp; GUARDRAILS CERTIFICATION", title_style))
    elements.append(Paragraph("Chronos Narrative Engine - Law Enforcement Report Generation System", subtitle_style))
    elements.append(HRFlowable(width="100%", thickness=2, color=HexColor('#1E3A8A'), spaceAfter=12))

    for heading, items in COMPLIANCE_SECTIONS:
        elements.append(Paragraph(heading, section_style))
        for item in items:
            safe = item.replace('<', '&lt;').replace('>', '&gt;')
            elements.append(Paragraph(f"* {safe}", bullet_style))

    return elements
